from docx import Document
from docx.shared import Inches
from PIL import Image
import xlrd
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor
def to_wps():
    strings = ["每张图片的标题"]
    images=['每张图片的名称']
    # 保存在本地的图片
    doc = Document()
    p=doc.add_heading(level=0)
    # p = doc.add_paragraph(style='ListBullet')
    run = p.add_run('网络运营日报')
    # doc.add_paragraph(style='ListBullet')
    # run.bold = True#加粗
    run.font.name = u'Cambria'#设置字体名称
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')
    run.font.color.rgb = RGBColor(0, 0, 0)#设置字体颜色
    p.paragraph_format.alignment =WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('1 关键链路使用情况明细（一周）')
    data = xlrd.open_workbook(r"/usr/local/python_on/Daily_automation_linux/xunjian/巡检报告.xlsx")
    # 获取sheet,通过序号
    table = data.sheet_by_index(0)
    excel = doc.add_table(rows=8, cols=5,style = "Table Grid")

    for i in range(0, 8):
        excel.cell(i, 0).text = table.cell(i, 0).value
        excel.cell(i, 1).text = table.cell(i, 1).value
        excel.cell(i, 2).text = table.cell(i, 2).value
        excel.cell(i, 3).text = table.cell(i, 3).value
        excel.cell(i, 4).text = table.cell(i, 4).value

    doc.add_paragraph()#输出一个空格

    for (string,image) in zip(strings,images):
        doc.add_paragraph(string) # 添加文字
        try:
            doc.add_picture(image, width=Inches(6))  # 添加图, 设置宽度
        except Exception:
            jpg_ima = Image.open(image)  # 打开图片
            jpg_ima.save('0.jpg')  # 保存新的图片
            doc.add_picture(image, width=Inches(6))  # 添加图, 设置宽度
    # doc.save('D:\python_on\Daily_automation\Daily word\网络运营日报.doc')  # 保存路径
    # doc.save('D:\python_on\Daily_automation\Daily word\%d.网络运营日报.doc' % float(time.time() * 10000000))  # 保存路径
    doc.save('/usr/local/python_on/Daily_automation_linux/Daily_word/网络运营日报.doc')

if __name__ == '__main__':
      to_wps()